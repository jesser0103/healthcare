import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


def add_image(doc: Document, image_path: Path) -> None:
    # Keep images reasonably sized for typical Word page width.
    # (6.5 inches ~= content width with normal margins)
    try:
        doc.add_picture(str(image_path), width=Pt(6.5 * 72))
        doc.add_paragraph("")
    except Exception as exc:
        add_paragraph(doc, f"[Image failed to load: {image_path.name}] {exc}")


def _set_mono(run):
    run.font.name = "Consolas"
    run.font.size = Pt(10)


def _set_normal(run):
    run.font.size = Pt(11)


def add_heading(doc: Document, text: str, level: int) -> None:
    # python-docx supports heading levels 0-9
    level = max(0, min(9, level))
    doc.add_heading(text.strip(), level=level)


def add_paragraph(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    _set_normal(run)


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    _set_normal(run)


def add_codeblock(doc: Document, lines: list[str]) -> None:
    for line in lines:
        p = doc.add_paragraph()
        run = p.add_run(line.rstrip("\n"))
        _set_mono(run)
    # small spacer
    doc.add_paragraph("")


def md_to_docx(md_path: Path, out_path: Path) -> None:
    md = md_path.read_text(encoding="utf-8")

    doc = Document()

    # Title page: first heading becomes title
    lines = md.splitlines()

    in_code = False
    code_lines: list[str] = []

    heading_re = re.compile(r"^(#{1,6})\s+(.*)$")
    bullet_re = re.compile(r"^\s*[-*]\s+(.*)$")
    image_re = re.compile(r"^!\[[^\]]*\]\(([^)]+)\)\s*$")

    first_title_done = False

    for raw in lines:
        line = raw.rstrip("\n")

        if line.strip().startswith("```"):
            if not in_code:
                in_code = True
                code_lines = []
            else:
                in_code = False
                add_codeblock(doc, code_lines)
                code_lines = []
            continue

        if in_code:
            code_lines.append(line)
            continue

        if not line.strip():
            doc.add_paragraph("")
            continue

        m = heading_re.match(line)
        if m:
            level = len(m.group(1))
            text = m.group(2).strip()
            if not first_title_done and level == 1:
                # Use a centered title paragraph for the main title
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(text)
                run.bold = True
                run.font.size = Pt(20)
                first_title_done = True
                doc.add_paragraph("")
            else:
                add_heading(doc, text, level)
            continue

        b = bullet_re.match(line)
        if b:
            add_bullet(doc, b.group(1).strip())
            continue

        img = image_re.match(line)
        if img:
            rel = img.group(1).strip()
            # Resolve relative to repo root (same folder as markdown)
            img_path = (md_path.parent / rel).resolve() if not Path(rel).is_absolute() else Path(rel)
            if img_path.exists():
                add_image(doc, img_path)
            else:
                add_paragraph(doc, f"[Missing image: {rel}]")
            continue

        # Default paragraph
        add_paragraph(doc, line)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))


if __name__ == "__main__":
    repo_root = Path(__file__).resolve().parents[1]
    md_path = repo_root / "COLLEGE_PROJECT_REPORT_8PAGES.md"
    out_path = repo_root / "COLLEGE_PROJECT_REPORT_8PAGES_DIAGRAMS.docx"

    if not md_path.exists():
        raise SystemExit(f"Markdown not found: {md_path}")

    md_to_docx(md_path, out_path)
    print(f"Wrote: {out_path}")
